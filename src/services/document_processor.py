"""
Document Processing Service
Text extraction, indexing, search, and AI-powered analysis
"""

import asyncio
import logging
import mimetypes
import hashlib
from typing import Dict, Any, List, Optional
from datetime import datetime
import re

# Document processing libraries
try:
    import PyPDF2
    import docx
    from bs4 import BeautifulSoup
    import markdown
except ImportError:
    # Fallback for basic text processing
    PyPDF2 = None
    docx = None
    BeautifulSoup = None
    markdown = None

from ..core.database import get_async_session
from ..models.document import Document, DocumentExtract, DocumentAnalysis
from sqlalchemy import select, update, and_, or_, desc
from sqlalchemy.orm import selectinload

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Service for processing uploaded documents"""

    def __init__(self):
        self.supported_types = {
            'application/pdf': 'pdf',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
            'text/plain': 'txt',
            'text/markdown': 'md',
            'text/html': 'html'
        }

    async def process_document(self, document_id: str) -> Dict[str, Any]:
        """Process a document: extract text, create index, analyze content"""

        try:
            logger.info(f"📄 Processing document: {document_id}")

            async with get_async_session() as session:
                # Get document
                stmt = select(Document).where(Document.id == document_id)
                result = await session.execute(stmt)
                document = result.scalar_one_or_none()

                if not document:
                    return {"success": False, "error": "Document not found"}

                # Update status to processing
                document.status = "processing"
                await session.commit()

                try:
                    # Extract text content
                    extraction_result = await self._extract_text(document)

                    if not extraction_result["success"]:
                        document.status = "error"
                        await session.commit()
                        return extraction_result

                    # Store extracted content
                    extract = DocumentExtract(
                        document_id=document_id,
                        content=extraction_result["content"],
                        extraction_method=extraction_result["method"],
                        word_count=len(extraction_result["content"].split()),
                        processing_metadata=extraction_result.get("metadata", {})
                    )

                    session.add(extract)

                    # Update document status
                    document.status = "ready"
                    document.processed_at = datetime.utcnow()

                    await session.commit()

                    logger.info(f"✅ Document processed successfully: {document_id}")

                    return {
                        "success": True,
                        "document_id": document_id,
                        "content_length": len(extraction_result["content"]),
                        "word_count": extract.word_count,
                        "extraction_method": extraction_result["method"]
                    }

                except Exception as e:
                    document.status = "error"
                    await session.commit()
                    logger.error(f"Error processing document {document_id}: {e}")
                    return {"success": False, "error": str(e)}

        except Exception as e:
            logger.error(f"Document processing failed: {e}")
            return {"success": False, "error": str(e)}

    async def _extract_text(self, document: Document) -> Dict[str, Any]:
        """Extract text content from document based on file type"""

        try:
            # Determine file type
            mime_type = mimetypes.guess_type(document.filename)[0]
            if not mime_type or mime_type not in self.supported_types:
                return {"success": False, "error": f"Unsupported file type: {mime_type}"}

            file_type = self.supported_types[mime_type]

            # Read file content
            if not document.file_content:
                return {"success": False, "error": "No file content available"}

            if file_type == 'pdf':
                return await self._extract_pdf_text(document.file_content)
            elif file_type == 'docx':
                return await self._extract_docx_text(document.file_content)
            elif file_type == 'txt':
                return await self._extract_txt_text(document.file_content)
            elif file_type == 'md':
                return await self._extract_markdown_text(document.file_content)
            elif file_type == 'html':
                return await self._extract_html_text(document.file_content)
            else:
                return {"success": False, "error": f"No extractor for file type: {file_type}"}

        except Exception as e:
            logger.error(f"Text extraction failed: {e}")
            return {"success": False, "error": str(e)}

    async def _extract_pdf_text(self, file_content: bytes) -> Dict[str, Any]:
        """Extract text from PDF file"""

        if not PyPDF2:
            return {"success": False, "error": "PDF processing not available (PyPDF2 required)"}

        try:
            import io
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))

            text_content = []
            metadata = {
                "page_count": len(pdf_reader.pages),
                "pdf_info": {}
            }

            # Extract PDF metadata if available
            if pdf_reader.metadata:
                metadata["pdf_info"] = {
                    "title": pdf_reader.metadata.get("/Title", ""),
                    "author": pdf_reader.metadata.get("/Author", ""),
                    "subject": pdf_reader.metadata.get("/Subject", ""),
                    "creator": pdf_reader.metadata.get("/Creator", "")
                }

            # Extract text from each page
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_content.append(f"[Page {page_num + 1}]\n{page_text}\n")
                except Exception as e:
                    logger.warning(f"Failed to extract text from page {page_num + 1}: {e}")
                    continue

            full_text = "\n".join(text_content)

            # Clean up text
            full_text = self._clean_extracted_text(full_text)

            return {
                "success": True,
                "content": full_text,
                "method": "PyPDF2",
                "metadata": metadata
            }

        except Exception as e:
            logger.error(f"PDF extraction failed: {e}")
            return {"success": False, "error": f"PDF extraction failed: {str(e)}"}

    async def _extract_docx_text(self, file_content: bytes) -> Dict[str, Any]:
        """Extract text from DOCX file"""

        if not docx:
            return {"success": False, "error": "DOCX processing not available (python-docx required)"}

        try:
            import io
            doc = docx.Document(io.BytesIO(file_content))

            text_content = []
            metadata = {
                "paragraph_count": 0,
                "table_count": 0
            }

            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
                    metadata["paragraph_count"] += 1

            # Extract text from tables
            for table in doc.tables:
                metadata["table_count"] += 1
                table_text = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        table_text.append(" | ".join(row_text))

                if table_text:
                    text_content.append("\n[TABLE]\n" + "\n".join(table_text) + "\n[/TABLE]\n")

            full_text = "\n".join(text_content)
            full_text = self._clean_extracted_text(full_text)

            return {
                "success": True,
                "content": full_text,
                "method": "python-docx",
                "metadata": metadata
            }

        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            return {"success": False, "error": f"DOCX extraction failed: {str(e)}"}

    async def _extract_txt_text(self, file_content: bytes) -> Dict[str, Any]:
        """Extract text from plain text file"""

        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']

            for encoding in encodings:
                try:
                    text = file_content.decode(encoding)
                    cleaned_text = self._clean_extracted_text(text)

                    return {
                        "success": True,
                        "content": cleaned_text,
                        "method": f"text_decode_{encoding}",
                        "metadata": {
                            "encoding": encoding,
                            "line_count": len(text.splitlines())
                        }
                    }
                except UnicodeDecodeError:
                    continue

            return {"success": False, "error": "Could not decode text file with any supported encoding"}

        except Exception as e:
            logger.error(f"Text extraction failed: {e}")
            return {"success": False, "error": f"Text extraction failed: {str(e)}"}

    async def _extract_markdown_text(self, file_content: bytes) -> Dict[str, Any]:
        """Extract text from Markdown file"""

        try:
            # Decode markdown content
            text = file_content.decode('utf-8')

            # If markdown library is available, convert to HTML first then extract text
            if markdown and BeautifulSoup:
                try:
                    html = markdown.markdown(text)
                    soup = BeautifulSoup(html, 'html.parser')
                    extracted_text = soup.get_text()
                    method = "markdown_to_html"
                except:
                    # Fallback to raw text
                    extracted_text = text
                    method = "markdown_raw"
            else:
                # Simple markdown text extraction (remove markdown syntax)
                extracted_text = self._simple_markdown_to_text(text)
                method = "markdown_simple"

            cleaned_text = self._clean_extracted_text(extracted_text)

            return {
                "success": True,
                "content": cleaned_text,
                "method": method,
                "metadata": {
                    "original_length": len(text),
                    "processed_length": len(cleaned_text)
                }
            }

        except Exception as e:
            logger.error(f"Markdown extraction failed: {e}")
            return {"success": False, "error": f"Markdown extraction failed: {str(e)}"}

    async def _extract_html_text(self, file_content: bytes) -> Dict[str, Any]:
        """Extract text from HTML file"""

        try:
            html = file_content.decode('utf-8')

            if BeautifulSoup:
                soup = BeautifulSoup(html, 'html.parser')
                # Remove script and style elements
                for script in soup(["script", "style"]):
                    script.decompose()
                text = soup.get_text()
                method = "beautifulsoup"
            else:
                # Simple HTML tag removal
                text = self._simple_html_to_text(html)
                method = "simple_html"

            cleaned_text = self._clean_extracted_text(text)

            return {
                "success": True,
                "content": cleaned_text,
                "method": method,
                "metadata": {
                    "original_html_length": len(html),
                    "extracted_text_length": len(cleaned_text)
                }
            }

        except Exception as e:
            logger.error(f"HTML extraction failed: {e}")
            return {"success": False, "error": f"HTML extraction failed: {str(e)}"}

    def _clean_extracted_text(self, text: str) -> str:
        """Clean and normalize extracted text"""

        # Remove excessive whitespace
        text = re.sub(r'\n\s*\n', '\n\n', text)  # Multiple newlines to double newline
        text = re.sub(r' +', ' ', text)  # Multiple spaces to single space
        text = re.sub(r'\t+', '\t', text)  # Multiple tabs to single tab

        # Remove non-printable characters (except newlines, tabs, and common punctuation)
        text = re.sub(r'[^\x20-\x7E\n\t]', '', text)

        # Trim whitespace
        text = text.strip()

        return text

    def _simple_markdown_to_text(self, markdown_text: str) -> str:
        """Simple markdown to text conversion without external libraries"""

        text = markdown_text

        # Remove headers
        text = re.sub(r'^#{1,6}\s*', '', text, flags=re.MULTILINE)

        # Remove bold and italic
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
        text = re.sub(r'\*(.*?)\*', r'\1', text)
        text = re.sub(r'__(.*?)__', r'\1', text)
        text = re.sub(r'_(.*?)_', r'\1', text)

        # Remove links
        text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)

        # Remove code blocks
        text = re.sub(r'```.*?```', '', text, flags=re.DOTALL)
        text = re.sub(r'`([^`]+)`', r'\1', text)

        return text

    def _simple_html_to_text(self, html: str) -> str:
        """Simple HTML to text conversion without external libraries"""

        # Remove script and style content
        text = re.sub(r'<script.*?</script>', '', html, flags=re.DOTALL | re.IGNORECASE)
        text = re.sub(r'<style.*?</style>', '', text, flags=re.DOTALL | re.IGNORECASE)

        # Remove HTML tags
        text = re.sub(r'<[^>]+>', '', text)

        # Decode HTML entities
        html_entities = {
            '&amp;': '&',
            '&lt;': '<',
            '&gt;': '>',
            '&quot;': '"',
            '&#39;': "'",
            '&nbsp;': ' '
        }

        for entity, char in html_entities.items():
            text = text.replace(entity, char)

        return text

    async def get_document_content(self, document_id: str) -> Dict[str, Any]:
        """Get processed content for a document"""

        try:
            async with get_async_session() as session:
                # Get document with extract
                stmt = select(Document).options(
                    selectinload(Document.extracts)
                ).where(Document.id == document_id)

                result = await session.execute(stmt)
                document = result.scalar_one_or_none()

                if not document:
                    return {"success": False, "error": "Document not found"}

                # Check if document is processed
                if document.status != "ready":
                    if document.status == "processing":
                        return {"success": False, "error": "Document is still being processed"}
                    elif document.status == "error":
                        return {"success": False, "error": "Document processing failed"}
                    else:
                        # Try to process if not started
                        process_result = await self.process_document(document_id)
                        if not process_result["success"]:
                            return process_result

                # Get latest extract
                if document.extracts:
                    latest_extract = max(document.extracts, key=lambda e: e.created_at)

                    return {
                        "success": True,
                        "content": latest_extract.content,
                        "metadata": {
                            "title": document.title,
                            "document_type": document.document_type,
                            "authors": document.authors,
                            "specialty": document.specialty,
                            "word_count": latest_extract.word_count,
                            "extraction_method": latest_extract.extraction_method,
                            "processing_timestamp": latest_extract.created_at.isoformat()
                        }
                    }
                else:
                    return {"success": False, "error": "No extracted content available"}

        except Exception as e:
            logger.error(f"Failed to get document content: {e}")
            return {"success": False, "error": str(e)}

    async def search_documents(
        self,
        query: str,
        document_ids: Optional[List[str]] = None,
        search_type: str = "semantic",
        max_results: int = 20,
        include_snippets: bool = True
    ) -> Dict[str, Any]:
        """Search through processed documents"""

        try:
            async with get_async_session() as session:
                # Build base query
                base_stmt = select(Document, DocumentExtract).join(
                    DocumentExtract, Document.id == DocumentExtract.document_id
                ).where(Document.status == "ready")

                # Filter by document IDs if provided
                if document_ids:
                    base_stmt = base_stmt.where(Document.id.in_(document_ids))

                # Execute query
                result = await session.execute(base_stmt)
                documents_with_extracts = result.all()

                # Perform search based on type
                if search_type == "keyword":
                    search_results = self._keyword_search(documents_with_extracts, query, max_results)
                elif search_type == "semantic":
                    # For now, fall back to keyword search
                    # In production, you would use vector embeddings and similarity search
                    search_results = self._keyword_search(documents_with_extracts, query, max_results)
                else:
                    search_results = self._keyword_search(documents_with_extracts, query, max_results)

                # Add snippets if requested
                if include_snippets:
                    for result in search_results:
                        result["snippet"] = self._generate_snippet(result["content"], query)

                return {
                    "success": True,
                    "results": search_results[:max_results],
                    "total_found": len(search_results),
                    "documents_searched": len(documents_with_extracts),
                    "search_time_ms": 0  # Would track actual search time
                }

        except Exception as e:
            logger.error(f"Document search failed: {e}")
            return {"success": False, "error": str(e)}

    def _keyword_search(self, documents_with_extracts, query: str, max_results: int) -> List[Dict[str, Any]]:
        """Perform keyword-based search"""

        results = []
        query_terms = query.lower().split()

        for document, extract in documents_with_extracts:
            content_lower = extract.content.lower()

            # Count term matches
            matches = 0
            for term in query_terms:
                matches += content_lower.count(term)

            if matches > 0:
                # Calculate relevance score
                relevance = matches / len(query_terms)

                results.append({
                    "document_id": document.id,
                    "title": document.title,
                    "document_type": document.document_type,
                    "authors": document.authors,
                    "specialty": document.specialty,
                    "content": extract.content,
                    "word_count": extract.word_count,
                    "relevance_score": relevance,
                    "matches": matches
                })

        # Sort by relevance
        results.sort(key=lambda x: x["relevance_score"], reverse=True)
        return results

    def _generate_snippet(self, content: str, query: str, snippet_length: int = 200) -> str:
        """Generate a snippet around query terms"""

        query_terms = query.lower().split()
        content_lower = content.lower()

        # Find first occurrence of any query term
        first_match_pos = None
        for term in query_terms:
            pos = content_lower.find(term)
            if pos != -1:
                if first_match_pos is None or pos < first_match_pos:
                    first_match_pos = pos

        if first_match_pos is None:
            return content[:snippet_length] + "..."

        # Create snippet around the match
        start = max(0, first_match_pos - snippet_length // 2)
        end = min(len(content), start + snippet_length)

        snippet = content[start:end]

        # Add ellipsis if truncated
        if start > 0:
            snippet = "..." + snippet
        if end < len(content):
            snippet = snippet + "..."

        return snippet

    async def store_analysis(
        self,
        document_id: str,
        analysis_type: str,
        analysis_content: str,
        metadata: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Store AI analysis results for a document"""

        try:
            async with get_async_session() as session:
                analysis = DocumentAnalysis(
                    document_id=document_id,
                    analysis_type=analysis_type,
                    analysis_content=analysis_content,
                    analysis_metadata=metadata
                )

                session.add(analysis)
                await session.commit()

                return {"success": True, "analysis_id": analysis.id}

        except Exception as e:
            logger.error(f"Failed to store analysis: {e}")
            return {"success": False, "error": str(e)}

    async def get_processing_status(self, document_id: str) -> Dict[str, Any]:
        """Get processing status for a document"""

        try:
            async with get_async_session() as session:
                stmt = select(Document).where(Document.id == document_id)
                result = await session.execute(stmt)
                document = result.scalar_one_or_none()

                if not document:
                    return {"success": False, "error": "Document not found"}

                return {
                    "success": True,
                    "document_id": document_id,
                    "status": document.status,
                    "uploaded_at": document.created_at.isoformat() if document.created_at else None,
                    "processed_at": document.processed_at.isoformat() if document.processed_at else None,
                    "file_size": document.file_size,
                    "file_type": mimetypes.guess_type(document.filename)[0]
                }

        except Exception as e:
            logger.error(f"Failed to get processing status: {e}")
            return {"success": False, "error": str(e)}

    async def reprocess_document(self, document_id: str) -> Dict[str, Any]:
        """Reprocess a document"""

        try:
            # Simply call process_document again
            result = await self.process_document(document_id)
            return result

        except Exception as e:
            logger.error(f"Document reprocessing failed: {e}")
            return {"success": False, "error": str(e)}

    async def get_processing_analytics(self) -> Dict[str, Any]:
        """Get analytics about document processing"""

        try:
            async with get_async_session() as session:
                # Get document counts by status
                status_counts_stmt = select(Document.status,
                                          asyncio.create_task(session.execute(
                                              select(Document).where(Document.status == Document.status)
                                          )).result().rowcount)

                # Simple analytics for now
                total_docs_stmt = select(Document)
                total_docs_result = await session.execute(total_docs_stmt)
                total_docs = len(total_docs_result.all())

                ready_docs_stmt = select(Document).where(Document.status == "ready")
                ready_docs_result = await session.execute(ready_docs_stmt)
                ready_docs = len(ready_docs_result.all())

                processing_docs_stmt = select(Document).where(Document.status == "processing")
                processing_docs_result = await session.execute(processing_docs_stmt)
                processing_docs = len(processing_docs_result.all())

                error_docs_stmt = select(Document).where(Document.status == "error")
                error_docs_result = await session.execute(error_docs_stmt)
                error_docs = len(error_docs_result.all())

                return {
                    "total_documents": total_docs,
                    "ready_documents": ready_docs,
                    "processing_documents": processing_docs,
                    "error_documents": error_docs,
                    "processing_success_rate": (ready_docs / total_docs * 100) if total_docs > 0 else 0,
                    "timestamp": datetime.utcnow().isoformat()
                }

        except Exception as e:
            logger.error(f"Failed to get processing analytics: {e}")
            return {"error": str(e)}

# Global document processor instance
document_processor = DocumentProcessor()